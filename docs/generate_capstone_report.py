from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "ScropIDS_Capstone_Report.docx"

PROJECT_TITLE = "ScropIDS"
PROJECT_SUBTITLE = (
    "Multi-Tenant Cloud-Native Intrusion Detection Platform "
    "with Scheduler-Driven Aggregation and LLM-Assisted Threat Triage"
)

PROJECT_META = {
    "project_term": "January-May 2026",
    "course_code": "CSE439",
    "group_number": "[Project Group Number]",
    "mentor_name": "[Mentor Name]",
    "mentor_designation": "[Designation]",
    "school": "School of Computer Science and Engineering",
    "university": "Lovely Professional University, Punjab",
    "academic_year": "2025-2026",
    "students": [
        ("[Student Name 1]", "[Registration Number 1]"),
        ("[Student Name 2]", "[Registration Number 2]"),
        ("[Student Name 3]", "[Registration Number 3]"),
    ],
}

VERIFICATION_NOTES = [
    "Backend unit tests executed successfully: ./backend/.venv/bin/python manage.py test apps.core",
    "Frontend production build executed successfully: npm run build",
    "Go agent build executed successfully: go build -o /tmp/scropids-agent-verify ./cmd/agent",
]


def style_paragraph(
    paragraph,
    *,
    alignment: int | None = None,
    line_spacing: float = 1.5,
    space_after: int = 6,
    space_before: int = 0,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for style_name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_centered_paragraph(document: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.2, space_after=4)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_blank_lines(document: Document, count: int = 1) -> None:
    for _ in range(count):
        document.add_paragraph("")


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    separate_text = OxmlElement("w:t")
    separate_text.text = "Right-click and update the table of contents in Word if it does not refresh automatically."
    fld_separate.append(separate_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    heading.style.font.name = "Times New Roman"
    style_paragraph(
        heading,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.2,
        space_before=8 if level == 1 else 4,
        space_after=4,
    )


def add_paragraphs(document: Document, *paragraphs: str) -> None:
    for paragraph_text in paragraphs:
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        paragraph = document.add_paragraph(paragraph_text)
        style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value
            for paragraph in row[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=8)
    run = paragraph.add_run(text.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_placeholder_box(document: Document, title: str, description: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.2, space_after=2)
    title_run = paragraph.add_run(title + "\n")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)
    body_run = paragraph.add_run(description)
    body_run.font.name = "Times New Roman"
    body_run.font.size = Pt(11)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def add_full_page_placeholder(document: Document, title: str, description: str) -> None:
    add_page_break(document)
    add_heading(document, title, 2)
    add_placeholder_box(
        document,
        "Insert Final Evidence Here",
        description
        + "\n\nSuggested content: full-page screenshot, command output, architectural diagram, or signed validation evidence.",
    )
    add_blank_lines(document, 12)


def build_cover_page(document: Document) -> None:
    add_blank_lines(document, 4)
    add_centered_paragraph(document, "CAPSTONE PROJECT REPORT", bold=True, size=18)
    add_centered_paragraph(document, f"(Project Term {PROJECT_META['project_term']})", size=12)
    add_blank_lines(document, 2)
    add_centered_paragraph(document, PROJECT_TITLE.upper(), bold=True, size=16)
    add_centered_paragraph(document, PROJECT_SUBTITLE, size=13)
    add_blank_lines(document, 2)
    add_centered_paragraph(document, "Submitted by", bold=True, size=12)
    add_blank_lines(document, 1)
    for name, reg_no in PROJECT_META["students"]:
        add_centered_paragraph(document, f"{name}    Registration Number: {reg_no}", size=12)
    add_blank_lines(document, 2)
    add_centered_paragraph(document, f"Project Group Number: {PROJECT_META['group_number']}", size=12)
    add_centered_paragraph(document, f"Course Code: {PROJECT_META['course_code']}", size=12)
    add_blank_lines(document, 2)
    add_centered_paragraph(document, "Under the Guidance of", bold=True, size=12)
    add_centered_paragraph(document, PROJECT_META["mentor_name"], size=12)
    add_centered_paragraph(document, PROJECT_META["school"], size=12)
    add_centered_paragraph(document, PROJECT_META["university"], size=12)
    add_centered_paragraph(document, f"Academic Year: {PROJECT_META['academic_year']}", size=12)


def build_front_matter(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "PAC Form", 1)
    add_placeholder_box(
        document,
        "Department Form Placeholder",
        "Insert the signed PAC form or evaluation sheet issued by the department before final submission.",
    )

    add_page_break(document)
    add_heading(document, "DECLARATION", 1)
    add_paragraphs(
        document,
        (
            f"We hereby declare that the project work entitled '{PROJECT_TITLE} - {PROJECT_SUBTITLE}' is an "
            "authentic record of our own work carried out as a requirement of the Capstone Project for the award "
            "of B.Tech degree in Computer Science / Data Science under the guidance of the project mentor named "
            "on the title page. All information furnished in this report is based on our own study, design, "
            "implementation, and testing. No part of this work has been submitted elsewhere for any other degree "
            "or diploma."
        ),
        "Student signatures and dates may be added below before final submission.",
    )
    for name, _ in PROJECT_META["students"]:
        document.add_paragraph(f"{name}: ____________________    Date: ____________________")

    add_page_break(document)
    add_heading(document, "CERTIFICATE", 1)
    add_paragraphs(
        document,
        (
            "This is to certify that the declaration statement made by the students for the project "
            f"'{PROJECT_TITLE} - {PROJECT_SUBTITLE}' is correct to the best of my knowledge and belief. "
            "The work has been completed under my guidance and supervision. The present report reflects the "
            "students' original investigation, implementation effort, and analysis, and is fit for submission "
            "towards the capstone project requirement."
        ),
        (
            f"Signature and Name of the Mentor: {PROJECT_META['mentor_name']}\n"
            f"Designation: {PROJECT_META['mentor_designation']}\n"
            f"{PROJECT_META['school']}\n"
            f"{PROJECT_META['university']}\n"
            "Date: ____________________"
        ),
    )

    add_page_break(document)
    add_heading(document, "ACKNOWLEDGEMENT", 1)
    add_paragraphs(
        document,
        (
            "The successful completion of this project would not have been possible without the guidance, "
            "support, and encouragement of many people. We express our sincere gratitude to our project mentor "
            "for continuous feedback, technical direction, and academic support throughout the capstone cycle."
        ),
        (
            "We also acknowledge the broader open-source communities behind Django, Django REST Framework, "
            "React, Vite, Tailwind CSS, PostgreSQL, Redis, Celery, Go, and Docker. Their documentation and "
            "tooling enabled rapid prototyping, repeatable deployment, and a strong engineering foundation for "
            "the project."
        ),
        (
            "Finally, we thank our classmates, peers, and family members for their encouragement during the "
            "design, implementation, testing, and documentation stages of this project."
        ),
    )

    add_page_break(document)
    add_heading(document, "TABLE OF CONTENTS", 1)
    add_toc(document)


def build_detailed_api_reference(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "14. DETAILED API REFERENCE", 1)
    add_paragraphs(
        document,
        (
            "This appendix expands the operational API view of ScropIDS. The platform uses a mix of session-backed "
            "dashboard APIs for human users and token-backed endpoints for agents. The backend is structured so that "
            "tenancy, authentication, telemetry ingest, configuration, and alert review can all be expressed through "
            "clear REST-style routes."
        ),
        (
            "The most important design principle in this API layer is scope awareness. Human users act inside an "
            "organization context, while agents act inside an authenticated device identity. This separation is what "
            "allows the same platform to safely serve multiple tenants without mixing their data."
        ),
    )
    add_table(
        document,
        ["Endpoint", "Method", "Auth", "Purpose"],
        [
            ["/api/v1/health/", "GET", "None", "Service and database health check"],
            ["/api/v1/auth/register/", "POST", "None", "Create user and initial organization"],
            ["/api/v1/auth/login/", "POST", "None", "Start session for dashboard user"],
            ["/api/v1/auth/me/", "GET", "Session", "Return user profile and organization memberships"],
            ["/api/v1/organizations/", "GET/POST", "Session", "List or create organizations"],
            ["/api/v1/agent-enrollment-tokens/", "GET/POST", "Session", "Issue and review one-time enrollment tokens"],
            ["/api/v1/agents/access-token/", "GET/POST", "Session", "Read or rotate organization access token"],
            ["/api/v1/agents/enroll/", "POST", "Enrollment token", "Register agent using legacy one-time token"],
            ["/api/v1/agents/quick-enroll/", "POST", "Organization access token", "Register agent using org token"],
            ["/api/v1/ingest/events/", "POST", "Agent token", "Submit normalized event batches"],
            ["/api/v1/ingest/heartbeat/", "POST", "Agent token", "Refresh agent last-seen timestamp"],
            ["/api/v1/ingest/config/", "GET", "Agent token", "Fetch runtime collector and interval profile"],
            ["/api/v1/dashboard/overview/", "GET", "Session", "Return tenant metrics and charts"],
            ["/api/v1/alerts/", "GET/PATCH", "Session", "List alerts or update alert status"],
            ["/api/v1/llm-providers/", "GET/POST/PATCH", "Session", "Manage tenant LLM providers"],
            ["/api/v1/scheduler-configs/", "GET/PATCH", "Session", "Manage server and agent schedule"],
            ["/api/v1/agent-downloads/", "GET", "Session", "Return downloadable agent artifact manifest"],
            ["/api/v1/agents/<uuid>/timeline/", "GET", "Session", "Return event and LLM history for one agent"],
        ],
    )

    add_heading(document, "14.1 Tenant and Session Endpoints", 2)
    add_paragraphs(
        document,
        (
            "The authentication flow for dashboard users begins either with self-registration or with an existing "
            "session login. Once the user is authenticated, the backend returns the list of organizations the user "
            "belongs to. The frontend stores the active organization slug and includes it when tenant context must "
            "be selected explicitly."
        ),
        (
            "This mechanism is reflected in the React hook that exposes current user context to the rest of the "
            "application. It is one of the key building blocks that keeps the dashboard responsive while still "
            "respecting multi-tenant boundaries."
        ),
    )
    add_code_block(
        document,
        """
POST /api/v1/auth/register/
{
  "username": "analyst1",
  "password": "strong-password",
  "organization_name": "Acme SOC"
}

GET /api/v1/auth/me/
Response includes username, staff status, and all accessible organizations.
        """,
    )

    add_heading(document, "14.2 Agent Onboarding Endpoints", 2)
    add_paragraphs(
        document,
        (
            "ScropIDS supports two onboarding patterns. The compatibility path uses one-time enrollment tokens, "
            "while the recommended operational path uses an organization access token. The second mode is better "
            "suited for real deployments because it allows administrators to rotate a single tenant token and reuse "
            "it during a device rollout campaign."
        ),
        (
            "Once onboarding succeeds, the backend returns an agent_id and agent_token pair. These credentials are "
            "then used for all telemetry and heartbeat communication."
        ),
    )
    add_code_block(
        document,
        """
POST /api/v1/agents/quick-enroll/
{
  "organization_slug": "acme-soc",
  "access_token": "<ORG_ACCESS_TOKEN>",
  "hostname": "win-lab-01",
  "operating_system": "windows",
  "ip_address": "203.0.113.10"
}

Response:
{
  "agent_id": "uuid",
  "agent_token": "opaque-agent-token",
  "organization_slug": "acme-soc"
}
        """,
    )

    add_heading(document, "14.3 Telemetry Ingest Endpoints", 2)
    add_paragraphs(
        document,
        (
            "Telemetry ingest is intentionally strict. The backend expects a well-formed JSON document containing an "
            "events array. Each event includes timestamp, event_type, severity, and a nested data payload. This "
            "structure keeps collection modules flexible while still preserving a stable top-level schema."
        ),
        (
            "The ingest endpoint also filters scaffold events that exist only for demonstration purposes. This "
            "keeps test data from polluting the aggregate pipeline when the agent starter emits baseline examples."
        ),
    )
    add_code_block(
        document,
        """
POST /api/v1/ingest/events/
Headers:
  X-Agent-ID: <uuid>
  X-Agent-Token: <token>

Body:
{
  "events": [
    {
      "timestamp": "2026-02-24T19:12:22Z",
      "event_type": "process_creation",
      "severity": "high",
      "data": {
        "process_name": "powershell.exe",
        "command_line": "powershell -enc aW52b2tl",
        "parent_process": "explorer.exe",
        "user": "admin"
      }
    }
  ]
}
        """,
    )

    add_heading(document, "14.4 Configuration, Alert, and Artifact Endpoints", 2)
    add_paragraphs(
        document,
        (
            "The configuration surface of ScropIDS is intentionally broad because the system is designed to be "
            "operated as a tenant-facing platform. Scheduler settings, rule packs, LLM providers, and agent "
            "download manifests are all treated as first-class operational resources instead of hidden backend knobs."
        ),
        (
            "This approach makes the reportable workflow much stronger: what administrators see in the report is the "
            "same set of controls they exercise in the application."
        ),
    )
    add_table(
        document,
        ["Dashboard Feature", "Primary Endpoint(s)", "Operational Benefit"],
        [
            ["Overview page", "/dashboard/overview/, /alerts/", "Fast SOC-style visibility into tenant health"],
            ["Alerts page", "/alerts/", "Status tracking, review, and decision support"],
            ["Agents page", "/agents/, /agents/access-token/, /agent-downloads/", "Onboarding and packaging workflow"],
            ["LLM Config page", "/llm-providers/", "Tenant-specific inference configuration"],
            ["Scheduler page", "/scheduler-configs/", "Controls both backend windows and agent runtime"],
            ["Rules page", "/scheduler-configs/ + rule_config_json", "Custom escalation behavior per tenant"],
        ],
    )


def build_database_dictionary(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "15. DATABASE DICTIONARY AND DATA DESIGN", 1)
    add_paragraphs(
        document,
        (
            "The database layer is central to the reliability of ScropIDS. Every key record is organization-aware, "
            "which means the tenancy boundary is embedded into the schema instead of being managed only at the UI "
            "layer. This reduces the risk of accidental cross-tenant access and simplifies query scoping."
        ),
        (
            "The event and aggregate models also support JSON-centric storage. This allows the platform to store "
            "normalized but extensible event payloads while still indexing the high-value columns needed for search, "
            "aggregation, and dashboard summaries."
        ),
    )
    add_heading(document, "15.1 Core Tenant Entities", 2)
    add_table(
        document,
        ["Model", "Field", "Meaning"],
        [
            ["Organization", "id", "Primary tenant identifier (UUID)"],
            ["Organization", "name", "Human-readable tenant name"],
            ["Organization", "slug", "Stable tenant key used in routing and headers"],
            ["Organization", "created_by", "User who created the organization"],
            ["Organization", "agent_access_token_encrypted", "Encrypted tenant token for quick enrollment"],
            ["Organization", "agent_access_token_rotated_at", "Timestamp of last token rotation"],
            ["OrganizationMembership", "organization", "Tenant referenced by the membership"],
            ["OrganizationMembership", "user", "User granted access"],
            ["OrganizationMembership", "role", "OWNER, ADMIN, ANALYST, or VIEWER"],
        ],
    )

    add_heading(document, "15.2 Agent and Enrollment Tables", 2)
    add_table(
        document,
        ["Model", "Field", "Meaning"],
        [
            ["AgentEnrollmentToken", "token_hash", "SHA-256 hash of the raw one-time token"],
            ["AgentEnrollmentToken", "description", "Admin note about the rollout wave or use case"],
            ["AgentEnrollmentToken", "expires_at", "Expiration time for safe onboarding"],
            ["AgentEnrollmentToken", "used_at", "Marks token as single-use once consumed"],
            ["Agent", "id", "Primary device identifier (UUID)"],
            ["Agent", "hostname", "Endpoint hostname shown in the dashboard"],
            ["Agent", "operating_system", "Operating system reported during enrollment"],
            ["Agent", "ip_address", "Optional device IP address"],
            ["Agent", "api_token_hash", "Stored hash of raw device API token"],
            ["Agent", "last_seen", "Last successful heartbeat or event ingest time"],
        ],
    )

    add_heading(document, "15.3 Event and Aggregate Tables", 2)
    add_table(
        document,
        ["Model", "Field", "Meaning"],
        [
            ["Event", "event_type", "Normalized label such as process_creation or network_snapshot"],
            ["Event", "severity", "Low / Medium / High / Critical source severity"],
            ["Event", "raw_json", "Flexible JSON payload carrying event details"],
            ["Event", "source_timestamp", "Time generated at the source"],
            ["Event", "processed", "Whether the record has been absorbed into an aggregate"],
            ["AggregatedWindow", "window_start", "Start timestamp for aggregation interval"],
            ["AggregatedWindow", "window_end", "End timestamp for aggregation interval"],
            ["AggregatedWindow", "source_event_count", "How many source events contributed"],
            ["AggregatedWindow", "summary_json", "Derived counters and dominant signals"],
            ["AggregatedWindow", "analyzed", "Whether the LLM/rule stage has completed"],
            ["AggregatedWindow", "llm_output", "Validated analysis JSON and escalation metadata"],
        ],
    )

    add_heading(document, "15.4 Alert, Scheduler, and Provider Tables", 2)
    add_table(
        document,
        ["Model", "Field", "Meaning"],
        [
            ["Alert", "threat_level", "Final low/medium/high/critical classification"],
            ["Alert", "confidence", "0-100 confidence value after analysis and escalation"],
            ["Alert", "llm_analysis", "Reasoning, action, provider, and related JSON context"],
            ["Alert", "status", "open, in_progress, or resolved"],
            ["SchedulerConfig", "aggregation_interval", "Seconds between backend aggregation windows"],
            ["SchedulerConfig", "min_severity", "Lowest event severity admitted to aggregation"],
            ["SchedulerConfig", "alert_threshold", "Threat level required before alert creation"],
            ["SchedulerConfig", "agent_sync_interval", "How often agents refresh runtime settings"],
            ["SchedulerConfig", "agent_event_interval", "How often agents send event batches"],
            ["SchedulerConfig", "rule_config_json", "Built-in and custom rule definitions"],
            ["LLMProviderConfig", "provider_type", "openai_compatible or ollama"],
            ["LLMProviderConfig", "base_url", "Remote or local model endpoint"],
            ["LLMProviderConfig", "model", "Concrete model name selected for analysis"],
            ["LLMProviderConfig", "encrypted_api_key", "Encrypted provider credential"],
        ],
    )
    add_paragraphs(
        document,
        (
            "The indexing choices in the event and aggregate tables are aligned with operational usage. Queries most "
            "often filter by organization, processed state, created_at, and agent. This means the schema is already "
            "optimized for the workflows that matter most: aggregation windows, alert retrieval, and dashboard counts."
        ),
        (
            "From a report perspective, this data model is one of the strongest parts of the project because it "
            "demonstrates that tenancy, operational scheduling, and explainable analysis were considered together "
            "instead of as isolated features."
        ),
    )


def build_module_catalogue(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "16. DETAILED MODULE CATALOGUE", 1)
    add_paragraphs(
        document,
        (
            "This appendix maps the repository structure to functional responsibility. It is useful during viva, "
            "review meetings, and future maintenance because it shows which code areas are responsible for which "
            "capabilities in the platform."
        ),
    )
    add_heading(document, "16.1 Backend Service Catalogue", 2)
    add_table(
        document,
        ["File / Module", "Responsibility", "Why It Matters"],
        [
            ["backend/apps/core/models.py", "Core tenant, agent, event, schedule, aggregate, alert, and provider models", "Defines the system data contract"],
            ["backend/apps/core/views.py", "REST endpoints for auth, onboarding, ingest, dashboard, and admin workflows", "Exposes the full application surface"],
            ["backend/apps/core/services/pipeline.py", "Aggregation windows, rule escalation, and LLM execution flow", "Implements the central triage pipeline"],
            ["backend/apps/core/services/llm.py", "Provider selection, outbound model calls, and strict JSON validation", "Prevents malformed reasoning from polluting alerts"],
            ["backend/apps/core/services/alerts.py", "Alert creation threshold logic", "Turns analysis into actionable queue items"],
            ["backend/apps/core/services/encryption.py", "Secret encryption and decryption", "Protects provider API keys and tokens at rest"],
            ["backend/apps/core/tenancy.py", "Organization resolution and role enforcement", "Maintains tenant isolation"],
            ["backend/apps/core/authentication.py", "Agent header authentication", "Secures ingest and heartbeat requests"],
        ],
    )

    add_heading(document, "16.2 Frontend Page Catalogue", 2)
    add_table(
        document,
        ["Page", "Primary Goal", "Important UI Behavior"],
        [
            ["DashboardPage.tsx", "Show live SOC-style overview", "Refreshes overview and alerts on interval, supports alert sound and quick drill-down"],
            ["AlertsPage.tsx", "Manage and review alert queue", "Supports severity filtering, sound controls, drawer details, and status updates"],
            ["AgentsPage.tsx", "Onboard endpoints and download artifacts", "Resets organization token, suggests best package, and shows quick run commands"],
            ["AgentTimelinePage.tsx", "Inspect per-agent history", "Combines raw event timeline and LLM insight windows"],
            ["LlmConfigPage.tsx", "Manage model providers", "Supports presets for OpenAI-compatible and local endpoints"],
            ["SchedulerConfigPage.tsx", "Tune server and agent schedules", "Single profile drives both aggregation and collector behavior"],
            ["RulesPage.tsx", "Manage escalation rules", "Edits built-in rules and imports/exports custom rule packs"],
            ["LoginPage.tsx", "Handle session sign-in", "Entry point for dashboard users"],
        ],
    )

    add_heading(document, "16.3 Agent Capability Catalogue", 2)
    add_paragraphs(
        document,
        (
            "The Go agent is intentionally designed as a starter that can grow into a richer endpoint component over "
            "time. Even in its current form it already demonstrates several important platform ideas: enrollment, "
            "local configuration persistence, runtime schedule synchronization, periodic telemetry submission, and "
            "cross-platform packaging."
        ),
    )
    add_table(
        document,
        ["Agent Capability", "Current Status", "Notes"],
        [
            ["Interactive setup wizard", "Implemented", "Supports organization token mode, direct credentials, and legacy enrollment token flow"],
            ["Local configuration file", "Implemented", "Stores API base, tenant slug, agent credentials, and host details"],
            ["Runtime config sync", "Implemented", "Fetches scheduler profile from /api/v1/ingest/config/"],
            ["Heartbeat messaging", "Implemented", "Updates device last_seen on the backend"],
            ["Process snapshot collection", "Implemented", "Captures OS-specific process listing"],
            ["Security log sampling", "Implemented", "Samples available system/auth logs where accessible"],
            ["Network snapshot collection", "Implemented", "Samples connection list and derives count hints"],
            ["File integrity monitoring", "Planned placeholder", "Structure exists but deep collector is future work"],
        ],
    )

    add_heading(document, "16.4 Deployment Service Catalogue", 2)
    add_table(
        document,
        ["Docker Service", "Purpose", "Key Runtime Notes"],
        [
            ["postgres", "Primary relational database", "Stores tenant state, events, aggregates, alerts, and configuration"],
            ["redis", "Broker and transient coordination", "Used by Celery worker and beat"],
            ["backend", "Django API and main application", "Runs migrations, collectstatic, and gunicorn"],
            ["worker", "Celery async worker", "Processes scheduled and queued jobs"],
            ["beat", "Celery beat scheduler", "Triggers periodic scheduler tick"],
            ["frontend", "Serves compiled React application", "Exposes HTTP and HTTPS ports for operator access"],
        ],
    )
    add_paragraphs(
        document,
        (
            "This modular breakdown makes the project easier to defend academically because every major system claim "
            "maps cleanly to code. It also gives future contributors a direct roadmap for where to add new collectors, "
            "visualizations, or security controls without destabilizing the rest of the platform."
        ),
    )


def build_validation_appendix(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "17. VALIDATION EVIDENCE AND TEST CASES", 1)
    add_paragraphs(
        document,
        (
            "The capstone evaluation for ScropIDS is strongest when the report demonstrates not just design intent "
            "but repeatable validation. This appendix therefore combines automated checks, smoke commands, and "
            "manual operator test cases that can be replayed during demonstration."
        ),
    )
    add_table(
        document,
        ["Test ID", "Scenario", "Expected Result"],
        [
            ["TC-01", "Register user and create first tenant", "Session is created and organization appears in /auth/me/"],
            ["TC-02", "Create organization through organizations endpoint", "New tenant slug is returned"],
            ["TC-03", "Issue one-time enrollment token", "Token record is created with expiry metadata"],
            ["TC-04", "Rotate organization access token", "Fresh tenant token is returned and previous token becomes obsolete"],
            ["TC-05", "Quick-enroll agent using organization token", "agent_id and agent_token are returned"],
            ["TC-06", "Send valid heartbeat", "Backend returns ok and updates last_seen"],
            ["TC-07", "Send valid event batch", "accepted count increases and event records are stored"],
            ["TC-08", "Fetch runtime config from agent endpoint", "Collector toggles and intervals match tenant schedule"],
            ["TC-09", "Run scheduler tick", "Aggregate windows are created for eligible events"],
            ["TC-10", "Analyze pending aggregate with active provider", "LLM output is validated and stored"],
            ["TC-11", "Trigger rule escalation path", "Final threat level may increase above raw model output"],
            ["TC-12", "Cross alert threshold", "Alert appears in alert queue"],
            ["TC-13", "Open alert in dashboard drawer", "Reasoning and recommended action are shown"],
            ["TC-14", "Download agent artifact manifest", "Platform-specific packages appear in response"],
            ["TC-15", "Build frontend bundle", "Production assets are generated successfully"],
            ["TC-16", "Run backend unit tests", "All current tests pass"],
            ["TC-17", "Build Go agent", "Binary compilation completes without errors"],
        ],
    )

    add_heading(document, "17.1 Command Evidence", 2)
    add_code_block(
        document,
        """
Backend unit tests:
./backend/.venv/bin/python manage.py test apps.core

Frontend production build:
npm run build

Go agent build:
go build -o /tmp/scropids-agent-verify ./cmd/agent
        """,
    )

    add_heading(document, "17.2 SaaS Smoke Validation Flow", 2)
    add_code_block(
        document,
        """
1. Start services
   docker compose up --build

2. Create admin user
   docker compose exec backend python manage.py createsuperuser

3. Create organization
   curl -u admin:adminpass -X POST http://localhost:8000/api/v1/organizations/ \
     -H "Content-Type: application/json" \
     -d '{"name":"Acme SOC"}'

4. Create enrollment token or rotate organization access token
5. Enroll agent
6. Send event batch
7. Trigger scheduler tick
8. Check dashboard overview
        """,
    )

    add_heading(document, "17.3 Risk-Based Validation Discussion", 2)
    add_paragraphs(
        document,
        (
            "Validation in ScropIDS is not limited to whether an endpoint responds. The more meaningful question is "
            "whether the platform behaves safely under realistic failure conditions. For this reason, the implementation "
            "contains explicit fallback logic when no active LLM provider exists or when provider output cannot be "
            "parsed into the required JSON schema."
        ),
        (
            "This matters in production-like security workflows because analysts cannot afford silent failure. Even "
            "when external inference is unavailable, the scheduler and rule engine still provide structure that keeps "
            "the platform operational and reviewable."
        ),
    )


def build_snapshot_appendix(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "18. SYSTEM SNAPSHOT AND EVIDENCE PAGES", 1)
    add_paragraphs(
        document,
        (
            "The following pages are intentionally reserved for high-value project evidence. When final screenshots "
            "are inserted, this section becomes a powerful visual appendix for demonstration, viva review, and final "
            "submission. Each page below corresponds directly to an implemented screen, command result, or deployment "
            "artifact in the repository."
        ),
    )

    placeholders = [
        ("18.1 Screenshot Page - Login", "Insert the final login page showing session-based entry into ScropIDS."),
        ("18.2 Screenshot Page - Registration / First Tenant Creation", "Insert the registration or first-run onboarding view that creates a tenant workspace."),
        ("18.3 Screenshot Page - Overview Dashboard", "Insert the main dashboard hero section with tenant title, alert counts, and live status."),
        ("18.4 Screenshot Page - Threat Distribution and KPI Widgets", "Insert the chart section that shows severity counts, totals, and dashboard KPIs."),
        ("18.5 Screenshot Page - Alerts Queue", "Insert the alert list with threat level, confidence, and status columns."),
        ("18.6 Screenshot Page - Alert Detail Drawer", "Insert the slide-over or detail panel with reasoning and recommended action."),
        ("18.7 Screenshot Page - Agents Page", "Insert the agent inventory table with hostname, OS, and online/offline state."),
        ("18.8 Screenshot Page - Token Rotation Panel", "Insert the organization access token panel after a successful reset."),
        ("18.9 Screenshot Page - Agent Artifact Downloads", "Insert the artifact list showing platform, package type, and download options."),
        ("18.10 Screenshot Page - Quick Run Command", "Insert the command block that helps operators start the downloaded agent."),
        ("18.11 Screenshot Page - Agent Timeline", "Insert the per-agent history page showing events and LLM insights."),
        ("18.12 Screenshot Page - LLM Provider Overview", "Insert the provider listing page with active/inactive configuration records."),
        ("18.13 Screenshot Page - LLM Provider Create/Edit Form", "Insert the create or edit form for provider configuration."),
        ("18.14 Screenshot Page - Scheduler Server Controls", "Insert the scheduler view showing aggregation interval and thresholds."),
        ("18.15 Screenshot Page - Scheduler Agent Profile", "Insert the collector and runtime profile controls for agents."),
        ("18.16 Screenshot Page - Rules Page Built-in Rules", "Insert the rule engine configuration with built-in detections."),
        ("18.17 Screenshot Page - Rules Page Custom Rules", "Insert the custom JSON rule editor or import/export controls."),
        ("18.18 Screenshot Page - Admin Interface", "Insert the Django admin or tenant administration view."),
        ("18.19 Evidence Page - Health Endpoint", "Insert terminal or browser output for /api/v1/health/ showing healthy services."),
        ("18.20 Evidence Page - Smoke Test cURL Responses", "Insert one or more terminal captures for organization creation, enrollment, or dashboard overview."),
        ("18.21 Evidence Page - Backend Unit Test Output", "Insert terminal evidence for successful backend tests."),
        ("18.22 Evidence Page - Frontend Build Output", "Insert terminal evidence for successful Vite production build."),
        ("18.23 Evidence Page - Go Agent Build Output", "Insert terminal evidence for successful Go agent compilation."),
        ("18.24 Evidence Page - Docker Service Topology", "Insert docker compose ps output or a deployment diagram capturing running services."),
    ]

    for title, description in placeholders:
        add_full_page_placeholder(document, title, description)


def build_project_management_appendix(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "19. DEPLOYMENT, RISK, AND PROJECT MANAGEMENT APPENDIX", 1)
    add_paragraphs(
        document,
        (
            "Capstone evaluation is stronger when technical depth is paired with delivery discipline. This appendix "
            "summarizes deployment readiness, likely risks, and a practical milestone plan that can be used to explain "
            "how the project evolved from concept to working platform."
        ),
    )
    add_heading(document, "19.1 Deployment Checklist", 2)
    add_numbered(
        document,
        [
            "Prepare backend environment variables and database credentials.",
            "Run migrations and collect static assets during backend startup.",
            "Validate PostgreSQL and Redis health checks before exposing the application.",
            "Confirm Celery worker and beat processes are connected to the same broker.",
            "Create an administrative account and at least one organization.",
            "Configure an LLM provider or deliberately test the rule-only fallback path.",
            "Rotate organization access token before onboarding devices in a shared lab.",
            "Verify health endpoint, dashboard overview, and alert queue before the demonstration.",
        ],
    )

    add_heading(document, "19.2 Risk Register", 2)
    add_table(
        document,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Provider API outage or invalid model output", "Alert enrichment quality may fall", "Use strict validation and fallback handling"],
            ["Improper tenant scoping", "Cross-tenant data exposure", "Enforce organization ownership and role checks at the API layer"],
            ["Noisy event ingestion", "Excessive false positives", "Use severity filtering, aggregation windows, and rule tuning"],
            ["Weak onboarding hygiene", "Unauthorized agents may attempt enrollment", "Use expiring tokens and rotate organization access tokens"],
            ["Long-running queue backlog", "Delayed alert creation", "Scale worker concurrency and monitor scheduler cadence"],
            ["Collector permission gaps", "Reduced endpoint visibility", "Expose scheduler guidance and document elevated-permission requirements"],
            ["Frontend bundle growth", "Slower load or heavier client delivery", "Introduce code splitting and chunk tuning in later iterations"],
            ["Limited long-term benchmarking", "Unknown scale ceiling", "Plan future load testing and performance profiling"],
        ],
    )

    add_heading(document, "19.3 Milestone Timeline", 2)
    add_table(
        document,
        ["Week / Phase", "Major Work Item", "Deliverable"],
        [
            ["Week 1", "Problem framing and platform selection", "Capstone scope definition"],
            ["Week 2", "Architecture planning and tenant workflow design", "Initial system blueprint"],
            ["Week 3", "Data model and API contract drafting", "Core backend schema"],
            ["Week 4", "Authentication and tenancy foundation", "User/tenant scoped backend"],
            ["Week 5", "Agent onboarding endpoints", "Enrollment and quick-enroll flows"],
            ["Week 6", "Telemetry ingest pipeline", "Event acceptance and storage"],
            ["Week 7", "Scheduler and aggregate window design", "Background analysis foundation"],
            ["Week 8", "LLM integration and JSON validation", "Model-assisted reasoning path"],
            ["Week 9", "Alert lifecycle and dashboard metrics", "Usable analyst outputs"],
            ["Week 10", "Frontend operational pages", "Agents, alerts, scheduler, rules, and LLM config UI"],
            ["Week 11", "Packaging and smoke validation", "Cross-platform agent artifacts and tests"],
            ["Week 12", "Documentation and final report preparation", "Submission-ready capstone package"],
        ],
    )

    add_heading(document, "19.4 Final Delivery Notes", 2)
    add_paragraphs(
        document,
        (
            "This project is already positioned well for an academic defense because it demonstrates a complete "
            "software lifecycle: problem analysis, requirements, architecture, implementation, validation, and "
            "deployment planning. The extra appendices also make it easier to defend the practical depth of the work."
        ),
        (
            "If exact institutional formatting or a hard page-count target must be met, the easiest next step is to "
            "replace the reserved evidence pages with final screenshots and command captures. That will increase the "
            "visual richness of the report without weakening its technical relevance."
        ),
    )


def build_supplementary_evidence_reserve(document: Document) -> None:
    add_page_break(document)
    add_heading(document, "20. SUPPLEMENTARY TECHNICAL EVIDENCE RESERVE", 1)
    add_paragraphs(
        document,
        (
            "The pages in this section are reserved for additional high-value material that often strengthens final "
            "capstone submissions: raw API captures, exported JSON structures, diagrams, validation outputs, and "
            "configuration snapshots. They are intentionally project-specific so the report can be enriched further "
            "without changing its overall structure."
        ),
    )

    reserve_pages = [
        ("20.1 Supplementary Page - Database Schema Visualization", "Insert a diagram or export that visualizes the relationship between organizations, agents, events, aggregates, alerts, and providers."),
        ("20.2 Supplementary Page - Organization and Membership Records", "Insert admin or API evidence showing how organizations and organization memberships are stored and reviewed."),
        ("20.3 Supplementary Page - Enrollment Token Issue Evidence", "Insert a capture of enrollment token creation with description, expiry, and one-time token behavior."),
        ("20.4 Supplementary Page - Quick-Enroll Response Capture", "Insert the response returned after successful quick enrollment using the organization access token."),
        ("20.5 Supplementary Page - Runtime Config Response", "Insert the JSON response returned from /api/v1/ingest/config/ for an enrolled agent."),
        ("20.6 Supplementary Page - Event Ingest Acceptance", "Insert the accepted/dropped response and a sample event payload after telemetry submission."),
        ("20.7 Supplementary Page - Aggregate Summary JSON", "Insert an example aggregate window summary showing failed logins, suspicious commands, and external connections."),
        ("20.8 Supplementary Page - Validated LLM Output JSON", "Insert a model response after strict schema validation and escalation metadata."),
        ("20.9 Supplementary Page - Alert Record JSON", "Insert a stored alert representation with threat_level, confidence, llm_analysis, and status."),
        ("20.10 Supplementary Page - Agent Local Configuration File", "Insert the saved agent configuration file or a redacted example of its structure."),
        ("20.11 Supplementary Page - Rule Pack Export JSON", "Insert a rule pack export from the Rules page or backend response."),
        ("20.12 Supplementary Page - Scheduler Profile Export", "Insert scheduler-config JSON showing server and agent runtime settings."),
        ("20.13 Supplementary Page - Artifact Manifest JSON", "Insert agent-download manifest output listing platform, architecture, package type, size, and download path."),
        ("20.14 Supplementary Page - Final Deployment Proof", "Insert final docker compose or browser evidence demonstrating that the full stack is live and operational."),
    ]

    for title, description in reserve_pages:
        add_full_page_placeholder(document, title, description)


def build_main_report(document: Document) -> None:
    add_page_break(document)

    add_heading(document, "2. PROFILE OF THE PROBLEM, RATIONALE & SCOPE OF THE STUDY", 1)
    add_heading(document, "2.1 Profile of the Problem", 2)
    add_paragraphs(
        document,
        (
            "Modern organizations generate a high volume of endpoint telemetry from workstations, servers, and "
            "mobile devices. Although these logs can reveal signs of credential abuse, malicious execution, or "
            "suspicious outbound communication, the data often arrives in inconsistent formats and in volumes "
            "that are difficult for small security teams to interpret quickly."
        ),
        (
            "The challenge becomes more complex in a multi-tenant Software-as-a-Service environment. A shared "
            "platform must isolate tenant data, support safe agent onboarding, provide centralized administration, "
            "and still expose only the information each analyst is authorized to see. Traditional tools often solve "
            "only part of this problem, forcing teams to combine multiple products and manual workflows."
        ),
    )

    add_heading(document, "2.2 Rationale", 2)
    add_paragraphs(
        document,
        (
            "The rationale behind ScropIDS is to build a single platform that combines endpoint telemetry "
            "collection, secure ingestion, scheduler-driven aggregation, explainable threat triage, and a usable "
            "dashboard. Instead of asking analysts to read raw logs first, the platform transforms collected data "
            "into summarized windows and prioritizes threats using a hybrid rule-and-LLM workflow."
        ),
        (
            "This project is especially relevant in academic and startup-like environments where cost, "
            "customizability, and reproducibility matter. By relying primarily on open-source components and a "
            "modular design, the platform remains suitable both for capstone evaluation and for future production "
            "hardening."
        ),
    )

    add_heading(document, "2.3 Scope of the Study", 2)
    add_bullets(
        document,
        [
            "Cross-platform endpoint agent starter for Windows, Linux, and macOS.",
            "Django + Django REST Framework backend for multi-tenant ingestion, aggregation, alerting, and administration.",
            "Scheduler-controlled analysis pipeline using Celery and Redis.",
            "Dual-mode LLM integration that supports both OpenAI-compatible APIs and local Ollama deployments.",
            "React-based dashboard for overview metrics, alerts, agents, LLM configuration, scheduler settings, and rules.",
            "Docker Compose deployment model for local demonstration and repeatable evaluation.",
        ],
    )
    add_paragraphs(
        document,
        (
            "The present scope focuses on a production-oriented platform foundation. It does not attempt to replace "
            "full commercial EDR tooling, deep kernel-level telemetry, or enterprise-scale orchestration in the "
            "current capstone phase."
        ),
    )

    add_heading(document, "2.4 Problem Statement", 2)
    add_paragraphs(
        document,
        (
            "The problem addressed in this capstone is the design and implementation of a cloud-native, "
            "multi-tenant intrusion detection platform that can securely onboard endpoint agents, accept normalized "
            "event streams, aggregate suspicious activity, produce explainable threat assessments, and present "
            "actionable alerts through a centralized web interface."
        ),
    )

    add_heading(document, "3. EXISTING SYSTEM", 1)
    add_heading(document, "3.1 Introduction", 2)
    add_paragraphs(
        document,
        (
            "Existing security monitoring approaches generally fall into four groups: host-based intrusion "
            "detection systems, security information and event management platforms, commercial endpoint "
            "detection and response suites, and ad-hoc manual monitoring. Each provides useful capabilities, "
            "but each also carries trade-offs related to complexity, cost, customization, or explainability."
        ),
    )

    add_heading(document, "3.2 Existing Software Analysis", 2)
    add_heading(document, "3.2.1 Traditional HIDS Platforms (OSSEC / Wazuh)", 3)
    add_paragraphs(
        document,
        (
            "Traditional host-based IDS tools are effective for centralized rule-based monitoring and integrity "
            "checks. They are strong on baseline detection workflows, but multi-tenant SaaS isolation, flexible "
            "cloud-native dashboards, and custom LLM-assisted analyst reasoning typically require additional "
            "integration work."
        ),
    )
    add_heading(document, "3.2.2 SIEM Platforms", 3)
    add_paragraphs(
        document,
        (
            "SIEM platforms provide broad log ingestion, search, and correlation. They are powerful for mature "
            "security teams, but the operational overhead, infrastructure cost, and tuning requirements can be "
            "significant for smaller organizations or academic projects."
        ),
    )
    add_heading(document, "3.2.3 Commercial EDR Platforms", 3)
    add_paragraphs(
        document,
        (
            "Commercial EDR solutions provide rich telemetry and strong incident-response workflows. However, "
            "they are often closed platforms with pricing and deployment assumptions that are not ideal for a "
            "custom capstone system intended for learning, experimentation, and transparent extensibility."
        ),
    )
    add_heading(document, "3.2.4 Manual Monitoring and Script-Based Review", 3)
    add_paragraphs(
        document,
        (
            "Many small teams still rely on shell scripts, local log files, and periodic manual inspection. "
            "This approach has low entry cost but poor scalability, inconsistent data formats, and delayed "
            "response cycles when suspicious activity spans multiple hosts or event types."
        ),
    )

    add_heading(document, "3.3 Data Flow Diagram: Existing System", 2)
    add_paragraphs(
        document,
        (
            "A simplified data flow for a conventional manual or loosely integrated monitoring setup is shown "
            "below."
        ),
    )
    add_code_block(
        document,
        """
Endpoint or Server Logs
        ->
Separate Local Tools / Scripts
        ->
Raw Log Files / Search Outputs
        ->
Manual Analyst Review
        ->
Delayed Escalation or Ticket Creation
        """,
    )

    add_heading(document, "3.4 Innovations in the Proposed ScropIDS System", 2)
    add_bullets(
        document,
        [
            "Tenant-scoped data model and access control for SaaS-style isolation.",
            "Organization access token and one-time enrollment token workflows for controlled onboarding.",
            "Scheduler-driven event aggregation that reduces alert noise before triage.",
            "Hybrid threat assessment using built-in rules plus validated LLM JSON output.",
            "Cross-platform agent packaging for Windows, Linux, and macOS.",
            "Unified web console for alerts, agents, LLM configuration, scheduler control, and rules.",
        ],
    )

    add_heading(document, "4. PROBLEM ANALYSIS", 1)
    add_heading(document, "4.1 Product Definition", 2)
    add_paragraphs(
        document,
        (
            "ScropIDS is a multi-tenant cloud-native intrusion detection platform. It collects normalized endpoint "
            "telemetry from distributed agents, stores events in a tenant-aware backend, groups suspicious activity "
            "into aggregate windows, evaluates those windows through rules and optional LLM analysis, and exposes "
            "the results through a dashboard designed for analysts and administrators."
        ),
    )

    add_heading(document, "4.2 Feasibility Analysis", 2)
    add_heading(document, "4.2.1 Technical Feasibility", 3)
    add_paragraphs(
        document,
        (
            "The project is technically feasible because it uses proven technologies with strong ecosystem support: "
            "Django and DRF for APIs, PostgreSQL for persistent data, Redis and Celery for scheduled background "
            "processing, React for the frontend, and Go for the endpoint agent. The repository already contains a "
            "working backend, dashboard pages, packaging scripts, and documented API flows."
        ),
    )
    add_heading(document, "4.2.2 Economic Feasibility", 3)
    add_paragraphs(
        document,
        (
            "The core platform is economically feasible for academic and small-scale deployment because the primary "
            "stack is open source. Optional LLM inference can be directed either to paid API providers or to local "
            "Ollama-based models, giving the operator flexibility based on budget and infrastructure."
        ),
    )
    add_heading(document, "4.2.3 Operational Feasibility", 3)
    add_paragraphs(
        document,
        (
            "The system is operationally feasible because administrators can manage organizations, tokens, agent "
            "downloads, scheduler settings, and LLM providers from a central web interface. Docker Compose also "
            "simplifies reproducible deployment for demos and controlled environments."
        ),
    )
    add_heading(document, "4.2.4 Time Feasibility", 3)
    add_paragraphs(
        document,
        (
            "The implementation is suitable for capstone timelines because it is modular. Core milestones such as "
            "requirements capture, data model design, API development, agent onboarding, dashboard workflows, and "
            "documentation can be completed incrementally and validated independently."
        ),
    )
    add_heading(document, "4.2.5 Legal and Ethical Feasibility", 3)
    add_paragraphs(
        document,
        (
            "Endpoint monitoring must always be deployed with organizational authorization, clear user policies, "
            "and least-privilege data handling. ScropIDS addresses this by emphasizing token-based authentication, "
            "tenant isolation, encrypted provider secrets, and a human-in-the-loop approach in which LLM output is "
            "treated as advisory rather than autonomous enforcement."
        ),
    )

    add_heading(document, "4.3 Project Plan", 2)
    add_table(
        document,
        ["Phase", "Primary Activities", "Key Outcome"],
        [
            ["Phase 1", "Requirement study, problem analysis, literature and tool review", "Scope and goals finalized"],
            ["Phase 2", "Architecture design, tenancy model, API and data schema planning", "Core design approved"],
            ["Phase 3", "Backend implementation for organizations, agents, ingest, aggregation, alerts", "Working service layer"],
            ["Phase 4", "Frontend dashboard pages and operational workflows", "Usable analyst/admin interface"],
            ["Phase 5", "Go agent onboarding, packaging, and runtime sync", "Cross-platform starter agent flow"],
            ["Phase 6", "Testing, smoke validation, report writing, and documentation", "Capstone deliverables ready"],
        ],
    )

    add_heading(document, "5. SOFTWARE REQUIREMENT ANALYSIS", 1)
    add_heading(document, "5.1 Introduction", 2)
    add_paragraphs(
        document,
        (
            "The software requirement analysis for ScropIDS focuses on the needs of three main stakeholders: tenant "
            "owners or administrators, security analysts, and endpoint agents. The system must provide secure "
            "onboarding, clear multi-tenant boundaries, reliable event handling, configurable detection logic, and "
            "a practical monitoring interface."
        ),
    )

    add_heading(document, "5.2 General Description", 2)
    add_table(
        document,
        ["Actor", "Role in the System"],
        [
            ["Tenant Owner / Admin", "Creates organizations, manages tokens, configures providers, adjusts scheduler and rules"],
            ["Security Analyst", "Views alerts, reviews reasoning, changes alert status, monitors dashboard metrics"],
            ["Endpoint Agent", "Enrolls with backend, sends event batches and heartbeat, receives runtime profile"],
        ],
    )

    add_heading(document, "5.3 Specific Requirements", 2)
    add_heading(document, "5.3.1 Functional Requirements", 3)
    add_table(
        document,
        ["ID", "Requirement"],
        [
            ["FR-01", "The system shall create and manage tenant organizations."],
            ["FR-02", "The system shall support secure agent onboarding through one-time tokens or organization access tokens."],
            ["FR-03", "The system shall authenticate enrolled agents using agent ID and agent token headers."],
            ["FR-04", "The system shall ingest normalized event batches and store them in tenant-scoped records."],
            ["FR-05", "The scheduler shall aggregate unprocessed events using configurable intervals and severity thresholds."],
            ["FR-06", "The system shall analyze aggregate windows using a configured LLM provider and a rule escalation engine."],
            ["FR-07", "The system shall create alerts when the resulting threat level crosses the configured threshold."],
            ["FR-08", "The dashboard shall expose overview, alerts, agents, LLM configuration, scheduler, and rule management pages."],
            ["FR-09", "The platform shall provide downloadable agent artifacts for multiple operating systems and architectures."],
        ],
    )

    add_heading(document, "5.3.2 Non-Functional Requirements", 3)
    add_table(
        document,
        ["ID", "Requirement"],
        [
            ["NFR-01", "Tenant data must remain isolated through organization ownership and scoped access control."],
            ["NFR-02", "Provider API keys must be stored securely and not exposed in plaintext to the frontend."],
            ["NFR-03", "The system should remain deployable through containerized services for reproducibility."],
            ["NFR-04", "The codebase should remain modular to support maintenance and future feature additions."],
            ["NFR-05", "The platform should present explainable alerts with reasoning and recommended action."],
            ["NFR-06", "The system should support portability across local demos and future cloud deployment targets."],
        ],
    )

    add_heading(document, "6. DESIGN", 1)
    add_heading(document, "6.1 System Architecture", 2)
    add_paragraphs(
        document,
        (
            "The ScropIDS architecture follows a layered flow. Endpoint agents collect local process, security, "
            "network, and system information. The backend exposes secure ingest endpoints that validate and store "
            "events. Celery-backed scheduler logic then groups unprocessed events into aggregate windows, applies "
            "LLM analysis and rule escalation, and persists alerts for dashboard consumption."
        ),
    )
    add_code_block(
        document,
        """
Endpoint Agent
        ->
Secure Ingest API (Django REST)
        ->
PostgreSQL Event Store
        ->
Scheduler + Aggregation
        ->
Rule Engine + LLM Analysis
        ->
Alerts
        ->
React Dashboard
        """,
    )

    add_heading(document, "6.2 Design Notations", 2)
    add_heading(document, "6.2.1 Data Flow Diagrams (DFD)", 3)
    add_paragraphs(
        document,
        (
            "At the context level, ScropIDS receives events from agents and configuration actions from dashboard "
            "users. It outputs alerts, dashboards, downloadable artifacts, and runtime schedules. At the lower "
            "level, the major data stores are organizations, agents, events, aggregated windows, alerts, scheduler "
            "profiles, and LLM provider settings."
        ),
    )
    add_heading(document, "6.2.2 Entity Relationship Diagram", 3)
    add_table(
        document,
        ["Entity", "Key Relationship"],
        [
            ["Organization", "Parent tenant entity for memberships, agents, events, scheduler configs, providers, and alerts"],
            ["OrganizationMembership", "Maps a user to an organization and role"],
            ["Agent", "Belongs to one organization and produces many events"],
            ["Event", "Belongs to one agent and one organization"],
            ["AggregatedWindow", "Summarizes a time bucket of related events for one agent"],
            ["Alert", "Created from analyzed aggregate windows according to threshold policy"],
            ["LLMProviderConfig", "Stores tenant-specific provider name, endpoint, model, and encrypted key"],
            ["SchedulerConfig", "Stores runtime intervals, collector toggles, thresholds, and rule configuration"],
        ],
    )

    add_heading(document, "6.3 Detailed Module Design", 2)
    add_heading(document, "6.3.1 Agent Collection and Enrollment Module", 3)
    add_paragraphs(
        document,
        (
            "The Go agent supports interactive setup, quick enrollment using organization access tokens, and "
            "legacy enrollment using one-time tokens. After onboarding, the agent stores configuration locally, "
            "periodically pulls the runtime schedule from /api/v1/ingest/config/, sends heartbeat messages, and "
            "submits event batches to the backend."
        ),
    )
    add_heading(document, "6.3.2 Secure Ingestion and Tenancy Module", 3)
    add_paragraphs(
        document,
        (
            "The backend validates incoming event payloads, authenticates agents via request headers, and stores "
            "tenant-aware event records. Human users access tenant data through session-backed dashboard APIs, "
            "with optional organization selection using the X-Organization-Slug request header."
        ),
    )
    add_heading(document, "6.3.3 Aggregation, Rule, and LLM Analysis Module", 3)
    add_paragraphs(
        document,
        (
            "Unprocessed events are grouped into aggregate windows by organization and agent. Each window captures "
            "summary counters such as failed logins, suspicious commands, new processes, and external connections. "
            "The resulting summary is sent to the selected LLM provider and then passed through a rule escalation "
            "engine so that deterministic rules can raise the final threat level when necessary."
        ),
    )
    add_heading(document, "6.3.4 Dashboard and Administration Module", 3)
    add_paragraphs(
        document,
        (
            "The dashboard exposes pages for live overview metrics, alert queue management, agent download and "
            "token workflows, LLM provider management, scheduler tuning, and rule editing. This keeps both analyst "
            "operations and administrator controls inside one tenant-aware interface."
        ),
    )

    add_heading(document, "6.4 Flowcharts", 2)
    add_heading(document, "6.4.1 Agent Enrollment and Event Flow", 3)
    add_code_block(
        document,
        """
Create organization
        ->
Generate enrollment or access token
        ->
Run agent setup / quick enroll
        ->
Receive agent_id and agent_token
        ->
Send heartbeat and event batches
        ->
Backend stores tenant-scoped events
        """,
    )
    add_heading(document, "6.4.2 Alert Generation Flow", 3)
    add_code_block(
        document,
        """
Collect unprocessed events
        ->
Filter by minimum severity
        ->
Group by agent and time window
        ->
Build aggregate summary
        ->
LLM analysis
        ->
Rule escalation
        ->
Compare with alert threshold
        ->
Create alert and show in dashboard
        """,
    )

    add_heading(document, "6.5 Pseudocode", 2)
    add_heading(document, "6.5.1 Aggregation Algorithm", 3)
    add_code_block(
        document,
        """
for each active scheduler configuration:
    determine aggregation window
    fetch unprocessed events for organization
    discard events below configured severity threshold
    group remaining events by agent
    for each agent group:
        create aggregate summary
        save aggregate window
    mark grouped events as processed
    analyze pending aggregates
        """,
    )
    add_heading(document, "6.5.2 Rule Escalation Algorithm", 3)
    add_code_block(
        document,
        """
threat_level = llm_output.threat_level
confidence = llm_output.confidence
for each enabled built-in or custom rule:
    if summary counters satisfy the rule:
        threat_level = max(threat_level, rule.set_threat)
        confidence = max(confidence, rule.min_confidence)
record matching rules
return updated llm_output
        """,
    )

    add_heading(document, "7. TESTING", 1)
    add_heading(document, "7.1 Functional Testing", 2)
    add_table(
        document,
        ["Test Area", "Purpose", "Current Evidence"],
        [
            ["Backend unit tests", "Validate token handling and aggregate processing behavior", VERIFICATION_NOTES[0]],
            ["Frontend build", "Confirm dashboard code compiles to a production bundle", VERIFICATION_NOTES[1]],
            ["Go agent build", "Confirm agent source builds successfully on the current workspace", VERIFICATION_NOTES[2]],
            ["SaaS smoke workflow", "Validate organization creation, enrollment, ingestion, scheduler tick, and overview retrieval", "Documented in docs/smoke-test-saas.md for end-to-end execution"],
        ],
    )

    add_heading(document, "7.2 Structural Testing (White-Box Testing)", 2)
    add_paragraphs(
        document,
        (
            "White-box testing focuses on internal service behavior, especially authentication, event aggregation, "
            "and state transitions. The current backend test suite includes explicit checks for agent token round-trip "
            "validation and for the aggregate_events pipeline marking source events as processed after creating an "
            "aggregate window."
        ),
    )

    add_heading(document, "7.3 Levels of Testing", 2)
    add_heading(document, "7.3.1 Unit Testing", 3)
    add_paragraphs(
        document,
        (
            "Unit tests are present in backend/apps/core/tests.py. These tests verify model behavior and aggregation "
            "logic without requiring the full deployed system."
        ),
    )
    add_heading(document, "7.3.2 Integration Testing", 3)
    add_paragraphs(
        document,
        (
            "Integration testing is supported through the documented API flow that creates an organization, issues "
            "tokens, enrolls an agent, ingests events, runs the scheduler, and reads dashboard metrics. These steps "
            "exercise interactions between authentication, persistence, background processing, and API serialization."
        ),
    )
    add_heading(document, "7.3.3 System Testing", 3)
    add_paragraphs(
        document,
        (
            "System testing is performed by deploying the services through Docker Compose and validating the full "
            "workflow from browser UI to backend APIs and worker-driven alert generation. This level ensures that "
            "service topology, environment configuration, and user-facing features operate together as expected."
        ),
    )
    add_heading(document, "7.3.4 User Acceptance Testing (UAT)", 3)
    add_paragraphs(
        document,
        (
            "UAT for this project centers on analyst and administrator tasks: signing in, choosing the active tenant, "
            "reviewing overview metrics, opening alerts, reading LLM reasoning, downloading agent packages, and "
            "editing scheduler or rule settings. The focus is on clarity of workflow and usefulness of alert context."
        ),
    )
    add_heading(document, "7.3.5 Performance Testing", 3)
    add_paragraphs(
        document,
        (
            "The current capstone phase emphasizes functional and architectural validation rather than large-scale "
            "load benchmarking. However, the design already anticipates horizontal scale by separating ingestion, "
            "storage, and background workers. Formal stress testing and queue-depth benchmarking remain part of "
            "future enhancement work."
        ),
    )

    add_heading(document, "8. IMPLEMENTATION", 1)
    add_heading(document, "8.1 Implementation of the Project", 2)
    add_paragraphs(
        document,
        (
            "The implementation is organized as a monorepo. The backend is written in Django and Django REST "
            "Framework, the frontend is built with React + Vite + TypeScript, and the endpoint agent is written in "
            "Go. Supporting services include PostgreSQL for persistent storage, Redis for message brokering, and "
            "Celery worker/beat processes for scheduler execution."
        ),
        (
            "The backend provides models for organizations, memberships, agents, events, scheduler profiles, LLM "
            "providers, aggregate windows, and alerts. Service modules handle encryption, LLM calls, alert creation, "
            "notifications, and pipeline execution. The frontend consumes tenant-filtered APIs and renders pages for "
            "overview, alerts, agents, LLM configuration, scheduler management, and rules."
        ),
    )

    add_heading(document, "8.2 Conversion Plan", 2)
    add_numbered(
        document,
        [
            "Deploy the backend, database, Redis, worker, beat, and frontend services using Docker Compose.",
            "Create the initial tenant organization and administrator account.",
            "Configure an active LLM provider or keep the platform in rule-only fallback mode during early rollout.",
            "Distribute agent packages to selected endpoints and enroll them using organization access tokens.",
            "Validate event flow and alert creation in a pilot environment before wider onboarding.",
            "Refine scheduler thresholds, rule packs, and alert handling policies based on analyst feedback.",
        ],
    )

    add_heading(document, "8.3 Post-Implementation and Software Maintenance", 2)
    add_bullets(
        document,
        [
            "Rotate organization access tokens and provider keys periodically.",
            "Maintain database backups and retention policies for event and alert records.",
            "Update dependencies, base images, and agent binaries as new fixes are released.",
            "Tune thresholds and custom rules to reduce false positives over time.",
            "Expand observability and incident audit trails as production adoption grows.",
        ],
    )

    add_heading(document, "9. PROJECT LEGACY", 1)
    add_heading(document, "9.1 Current Status of the Project", 2)
    add_paragraphs(
        document,
        (
            "ScropIDS currently exists as a production-oriented capstone implementation with a functioning backend, "
            "dashboard, multi-tenant model, token-based onboarding, scheduler-driven aggregation, rule escalation, "
            "and LLM provider integration. Cross-platform packaging support for downloadable agents is also present."
        ),
    )
    add_heading(document, "9.2 Remaining Areas of Concern", 2)
    add_bullets(
        document,
        [
            "Full OS-specific telemetry collectors are still evolving beyond the current starter implementation.",
            "Role-based access control can be further refined for larger analyst teams.",
            "Formal load testing and long-duration operational observation are still needed.",
            "Advanced threat enrichment such as Sigma rule packs and ATT&CK mapping remains future work.",
        ],
    )
    add_heading(document, "9.3 Technical and Managerial Lessons Learnt", 2)
    add_heading(document, "9.3.1 Technical Lessons", 3)
    add_bullets(
        document,
        [
            "Strict JSON contracts simplify downstream processing and reduce ambiguity in alerts.",
            "Combining deterministic rules with LLM reasoning produces a more dependable triage workflow than relying on either alone.",
            "A tenant-aware data model must be designed early, not added later as an afterthought.",
            "Packaging and onboarding are as important for usability as the detection logic itself.",
        ],
    )
    add_heading(document, "9.3.2 Managerial Lessons", 3)
    add_bullets(
        document,
        [
            "Incremental milestones reduced complexity by allowing backend, frontend, and agent work to progress in parallel.",
            "Continuous documentation made the system easier to validate and explain during review stages.",
            "A clear scope boundary prevented the capstone from drifting into enterprise-scale EDR ambitions.",
            "Feedback from likely users is critical for making alert output understandable and actionable.",
        ],
    )

    add_heading(document, "10. USER MANUAL (HELP GUIDE)", 1)
    add_heading(document, "10.1 Introduction", 2)
    add_paragraphs(
        document,
        (
            "This section explains the normal operator workflow for ScropIDS. The platform supports session-based "
            "dashboard usage for human users and token-based communication for endpoint agents."
        ),
    )
    add_heading(document, "10.2 System Requirements", 2)
    add_table(
        document,
        ["Component", "Minimum Requirement"],
        [
            ["Backend host", "2 vCPU, 4 GB RAM"],
            ["Database", "PostgreSQL 13+ with JSONB support"],
            ["Queue", "Redis 6+"],
            ["Agent endpoint", "Outbound network access to API and permission to read available audit sources"],
            ["Frontend", "Modern Chromium, Firefox, or Safari browser"],
        ],
    )
    add_heading(document, "10.3 Dashboard User Guide", 2)
    add_heading(document, "10.3.1 Account Registration and Login", 3)
    add_numbered(
        document,
        [
            "Open the application and navigate to the login page.",
            "Use an existing username/password pair or register a new account if the deployment allows self-registration.",
            "After authentication, select the active organization from the tenant switcher in the header.",
        ],
    )
    add_heading(document, "10.3.2 Reviewing Overview and Alerts", 3)
    add_numbered(
        document,
        [
            "Open the Overview page to review total agents, recent events, active alerts, and threat distribution.",
            "Open the Alerts page to inspect alert severity, confidence, reasoning, and recommended actions.",
            "Update alert status as work progresses from open to in_progress or resolved.",
        ],
    )
    add_heading(document, "10.3.3 Managing Agents and Tokens", 3)
    add_numbered(
        document,
        [
            "Open the Agents page to review enrolled endpoints and online/offline state.",
            "Reset or copy the organization access token when onboarding new machines.",
            "Download the recommended package for Windows, Linux, or macOS and run the displayed command.",
        ],
    )
    add_heading(document, "10.3.4 Reading LLM Recommendations", 3)
    add_paragraphs(
        document,
        (
            "For each alert, the drawer panel includes reasoning text and recommended action produced by the "
            "analysis pipeline. Analysts should treat this guidance as decision support and verify the suggestion "
            "before taking containment or remediation action."
        ),
    )

    add_heading(document, "10.4 Administrator Guide", 2)
    add_heading(document, "10.4.1 Creating Organizations and Memberships", 3)
    add_paragraphs(
        document,
        (
            "Administrators can create organizations through the API or the administrative interface. Memberships "
            "determine which users are allowed to access tenant data and configuration controls."
        ),
    )
    add_heading(document, "10.4.2 Configuring LLM Providers", 3)
    add_paragraphs(
        document,
        (
            "The LLM Config page supports both API-hosted and local providers. Enter a provider name, base URL, "
            "model name, timeout, and API key if needed. Activate only the provider intended for the current tenant."
        ),
    )
    add_heading(document, "10.4.3 Managing Scheduler and Rules", 3)
    add_paragraphs(
        document,
        (
            "The Scheduler page controls aggregation interval, minimum severity, alert threshold, and agent runtime "
            "collector toggles. The Rules page enables built-in rule escalation and supports custom rule JSON or "
            "imported rule packs."
        ),
    )

    add_heading(document, "10.5 Agent Setup Guide", 2)
    add_numbered(
        document,
        [
            "Download the correct platform package from the Agents page.",
            "Run the binary with --setup or provide SCROPIDS_API_BASE, SCROPIDS_ORG_SLUG, and SCROPIDS_ORG_ACCESS_TOKEN.",
            "Allow the agent to quick-enroll and save its local configuration.",
            "Verify that heartbeat and event batches appear in the backend.",
        ],
    )

    add_heading(document, "10.6 Troubleshooting Guide", 2)
    add_bullets(
        document,
        [
            "If agent enrollment fails, verify the organization slug, token value, and API base URL.",
            "If no alerts appear, check whether an active LLM provider or suitable rule thresholds are configured.",
            "If the dashboard is empty for a tenant, confirm that the correct organization is selected in the header.",
            "If scheduler-driven alerts lag behind, inspect Celery worker/beat services and Redis connectivity.",
        ],
    )

    add_heading(document, "11. SYSTEM SNAPSHOTS", 1)
    add_heading(document, "11.1 System Snapshots", 2)
    add_paragraphs(
        document,
        (
            "The following snapshot inventory can be used when final live screenshots are inserted before submission. "
            "Each caption already matches the implemented modules in the repository."
        ),
    )
    add_table(
        document,
        ["Suggested Screenshot", "Caption / Description"],
        [
            ["Login Page", "Session-based authentication screen for dashboard users"],
            ["Overview Dashboard", "Tenant-level metrics, live alerts, and threat distribution"],
            ["Alerts Page", "Alert queue with threat level, confidence, reasoning, and recommended action"],
            ["Agents Page", "Organization token reset, package downloads, and enrolled agent list"],
            ["LLM Config Page", "Tenant-scoped provider setup for OpenAI-compatible or local models"],
            ["Scheduler Page", "Aggregation interval, collector toggles, and agent runtime profile"],
            ["Rules Page", "Built-in rule escalation settings and custom rule pack import/export"],
        ],
    )
    add_placeholder_box(
        document,
        "Submission Note",
        "Insert the final screenshots after starting the full stack locally or through Docker Compose so the report includes the latest UI state.",
    )

    add_heading(document, "12. FUTURE SCOPE AND ENHANCEMENTS", 1)
    add_bullets(
        document,
        [
            "Implement deeper OS-specific collectors for Windows event logs, Linux audit sources, and macOS unified logs.",
            "Add Sigma-compatible rule pack support, MITRE ATT&CK mapping, and threat-intelligence enrichment.",
            "Improve role-based access control, audit logging, and analyst collaboration workflows.",
            "Introduce large-scale performance testing and orchestration-ready deployment options such as Kubernetes.",
            "Expand guided remediation playbooks and richer dashboard visualizations for investigation workflows.",
        ],
    )

    add_heading(document, "13. CONCLUSION", 1)
    add_paragraphs(
        document,
        (
            "ScropIDS demonstrates that a capstone project can move beyond a conceptual IDS prototype and become a "
            "practical, multi-tenant security platform. The system combines secure agent onboarding, normalized "
            "telemetry ingestion, scheduler-driven aggregation, rule-assisted LLM triage, and a dashboard that makes "
            "the output easier for analysts to act upon."
        ),
        (
            "The current implementation already provides a strong foundation for academic evaluation and future "
            "engineering growth. With further collector maturity, broader testing, and richer threat enrichment, "
            "the platform can evolve into an even more capable operational monitoring solution."
        ),
    )

    build_detailed_api_reference(document)
    build_database_dictionary(document)
    build_module_catalogue(document)
    build_validation_appendix(document)
    build_snapshot_appendix(document)
    build_project_management_appendix(document)
    build_supplementary_evidence_reserve(document)

    add_page_break(document)
    add_heading(document, "BIBLIOGRAPHY", 1)
    add_numbered(
        document,
        [
            "ScropIDS project repository documentation: README.md, docs/architecture.md, docs/api-contract.md, docs/smoke-test-saas.md.",
            "Django Documentation. https://docs.djangoproject.com/",
            "React Documentation. https://react.dev/",
            "Celery Documentation. https://docs.celeryq.dev/",
            "PostgreSQL Documentation. https://www.postgresql.org/docs/",
            "Docker Documentation. https://docs.docker.com/",
            "MITRE ATT&CK. https://attack.mitre.org/",
            "NIST Cybersecurity Framework 2.0. https://www.nist.gov/cyberframework",
        ],
    )


def build_document() -> Document:
    document = Document()
    configure_document(document)
    build_cover_page(document)
    build_front_matter(document)
    build_main_report(document)
    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Created {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
